import sqlite3
import json
from datetime import datetime

from docx import Document
from create_database import create_connection
from openai import OpenAI
import aspose.words as aw
import os


def fetch_applicant_data(uuid, conn):
    tables = [
        "applicants", "education", "projects", "work_experience",
        "volunteer_work", "languages", "certifications",
        "awards", "skills", "personal_projects"
    ]
    data = {}

    for table in tables:
        data[table] = []
        query = f"SELECT * FROM {table} WHERE applicant_uuid = ?"
        cursor = conn.cursor()
        cursor.execute(query, (uuid,))
        rows = cursor.fetchall()

        columns = [column[0] for column in cursor.description]

        for row in rows:
            row_data = dict(zip(columns, row))
            data[table].append(row_data)

    contact_info_dict = json.loads(data['applicants'][0]['contact_info'])
    email = contact_info_dict['email']
    phone = contact_info_dict['phone']
    full_name = data['applicants'][0]['full_name']
    return data, email, phone, full_name

def generate_cover_letter(applicant_data, job_data):
    client = None
    try:
        client = OpenAI(api_key='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')
    except Exception as e:
        print(str(e))
    prompt = (f"""You are a professional recruiter, perfect for helping candidates get their dream jobs.
                 Given the following job description and applicant details, write a cover letter that 
                 highlights the applicants qualities needed for the job (without making anything up).
                 The letter should be written in the 1st person, and should only take the relevant data to the job.
                 Output only the letter, without any explanation or details.
                 <applicant_data>: {applicant_data}
                 </applicant_data>
                 <job_details>: {job_data}
                 </job_data>""")
    chat_completion = client.chat.completions.create(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-3.5-turbo"
    )
    ai_output = chat_completion.choices[0].message.content
    print(ai_output)
    return(ai_output)


def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for run in inline:
            run.text = run.text.replace(key, value)

def update_word_template(file_path, contact_tuple, cover_letter_text):
    # Load the Word document
    doc = Document(file_path)

    today_date = datetime.today().strftime('%d %b, %Y')

    # Dictionary to map placeholders to their respective data
    replacements = {
        "Full Name": contact_tuple[2],  # Assuming index 2 is Full Name
        "Phone Number": contact_tuple[1],  # Assuming index 1 is Phone Number
        "Email": contact_tuple[0],  # Assuming index 0 is Email
        "Text": cover_letter_text,  # Cover letter text to replace 'Text' placeholder
        "Date": today_date  # Placeholder for today's date
    }

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            replace_text_in_paragraph(paragraph, key, value)

    # Replace placeholders in table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        replace_text_in_paragraph(paragraph, key, value)

    # Save the modified document
    new_file_path = "temporary_files/modified_CL_template.docx"
    doc.save(new_file_path)
    return new_file_path


def main(job_details):
    database = "database.db"
    uuid = "UUID_PLACEHOLDER"
    conn = create_connection(database)

    if conn is not None:
        applicant_data, email, phone, full_name = fetch_applicant_data(uuid, conn)
        conn.close()
        data_tuple = (email, phone, full_name)
        applicant_data_json = json.dumps(applicant_data, indent=4)
        cover_letter_text = generate_cover_letter(applicant_data_json, job_details)

        updated_file_path = update_word_template('CL template.docx', data_tuple, cover_letter_text)

        print(f"Updated document saved to: {updated_file_path}")
        return updated_file_path
    else:
        print("Error! Cannot create the database connection.")

if __name__ == "__main__":
    job_details = "Your job details here"
    something = update_word_template('CL template.docx', ('my email', '123213', 'Ilai :)'), 'I am cool')
